from docx import Document
from docx.shared import Pt
import re
import os

def markdown_to_docx(md_path, docx_path):
    document = Document()
    
    # Check if MD file exists
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    # Title
    document.add_heading('BhashaLLM Project Report', 0)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_paragraph = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Parse basic Markdown
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            # Bullet point
            document.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):
            # Table row (Just adding as plain text for simplicity in this script)
            # A full markdown table parser is complex; this keeps data readable.
            p = document.add_paragraph(line)
            p.style = 'No Spacing' # Use monospaced look if possible, or just tight
        else:
            # Normal paragraph
            # Bold handling (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = document.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    document.save(docx_path)
    print(f"DOCX generated: {docx_path}")

if __name__ == "__main__":
    markdown_to_docx("BhashaLLM/FINAL_RESEARCH_REPORT.md", "BhashaLLM/BhashaLLM_Report.docx")
